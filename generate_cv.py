"""
generate_cv.py
Generates Academic_CV_Josephine_Amponsah_Baah.docx using python-docx.

Install dependency if needed:
    pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour constants ───────────────────────────────────────────────────────────
GREEN = RGBColor(34, 139, 34)
BLACK = RGBColor(0, 0, 0)
DARK_GREY = RGBColor(64, 64, 64)

# ── Helper utilities ───────────────────────────────────────────────────────────

def set_run_font(run, name="Calibri", size_pt=10.5, bold=False, italic=False,
                 color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=0):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bottom_border(paragraph):
    """Add a thin bottom border (rule) beneath a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "228B22")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_header(doc, title):
    """Bold, green, uppercase section header with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    set_run_font(run, size_pt=11, bold=True, color=GREEN)
    add_bottom_border(p)
    return p


def add_bullet(doc, text, indent_level=0, size_pt=10.5):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.2)
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt)
    return p


def add_experience_header(doc, title, period, bold_title=True):
    """Bold role title + right-aligned period on same line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)

    run_title = p.add_run(title)
    set_run_font(run_title, size_pt=10.5, bold=bold_title)

    # Tab + period right-aligned via tab stop
    tab_run = p.add_run("\t" + period)
    set_run_font(tab_run, size_pt=10.5, italic=True, color=DARK_GREY)

    # Right tab stop at 6.5 inches (within 1-inch margins on letter paper)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")   # 6.5" × 1440 twips/inch
    tabs.append(tab)
    pPr.append(tabs)
    return p


def add_italic_label(doc, text, size_pt=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt, italic=True, color=DARK_GREY)
    return p


def inline_bold_normal(paragraph, bold_text, normal_text, size_pt=10.5):
    r1 = paragraph.add_run(bold_text)
    set_run_font(r1, size_pt=size_pt, bold=True)
    r2 = paragraph.add_run(normal_text)
    set_run_font(r2, size_pt=size_pt)


# ── Document setup ─────────────────────────────────────────────────────────────

doc = Document()

# Margins – 1 inch all sides
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── NAME & CONTACT HEADER ─────────────────────────────────────────────────────

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_before = Pt(0)
name_p.paragraph_format.space_after = Pt(2)
name_run = name_p.add_run("Josephine Amponsah Baah")
set_run_font(name_run, size_pt=22, bold=True)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(3)
title_run = title_p.add_run("ML/AI Researcher")
set_run_font(title_run, size_pt=12, italic=True, color=GREEN)

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_before = Pt(0)
contact_p.paragraph_format.space_after = Pt(6)
contact_run = contact_p.add_run(
    "josephine.amponsah98@gmail.com  |  "
    "github.com/josephine-amponsah  |  "
    "linkedin.com/in/josephine-amponsah-baah"
)
set_run_font(contact_run, size_pt=10, color=DARK_GREY)

# ── SECTION 1: RESEARCH INTERESTS ────────────────────────────────────────────

add_section_header(doc, "Research Interests")

research_text = (
    "My research interests lie at the intersection of Natural Language Processing, "
    "Large Language Models, and deep learning for sequence modelling and representation "
    "learning. I am particularly interested in studying the robustness and bias of NLI/NLU "
    "systems — examining how pre-trained transformers exploit spurious dataset artifacts and "
    "developing debiasing strategies that improve generalisation in high-stakes settings. "
    "Beyond core NLP, I am drawn to applications in Financial AI: leveraging statistical "
    "and machine learning methods for credit risk modelling and fraud detection, and "
    "exploring domain-adapted LLMs (FinBERT-style architectures) for intelligent financial "
    "document processing, validation, and compliance. I am actively seeking research "
    "opportunities that bridge robust natural language understanding with real-world "
    "financial and safety-critical applications."
)

ri_p = doc.add_paragraph()
ri_p.paragraph_format.space_before = Pt(3)
ri_p.paragraph_format.space_after = Pt(4)
ri_run = ri_p.add_run(research_text)
set_run_font(ri_run, size_pt=10.5)

# ── SECTION 2: EDUCATION ──────────────────────────────────────────────────────

add_section_header(doc, "Education")

# MSc
add_experience_header(
    doc,
    "MSc. Data Science  —  University of Texas, Austin, USA",
    "Aug 2024 – April 2026",
)

courses_p = doc.add_paragraph()
courses_p.paragraph_format.space_before = Pt(1)
courses_p.paragraph_format.space_after = Pt(1)
inline_bold_normal(
    courses_p,
    "Relevant Courses: ",
    "Machine Learning, Data Structures & Algorithms, Deep Learning, "
    "Advanced Predictive Modelling (Time-Series), Natural Language Processing, "
    "Advanced Deep Learning, Probability & Statistics",
)

proj_label_p = doc.add_paragraph()
proj_label_p.paragraph_format.space_before = Pt(2)
proj_label_p.paragraph_format.space_after = Pt(1)
proj_label_run = proj_label_p.add_run("Relevant Projects (from coursework):")
set_run_font(proj_label_run, size_pt=10.5, bold=True)

add_bullet(
    doc,
    "Optimized natural language understanding of BERT models with ensemble debiasing "
    "and data cartography (NLP Final Project)",
)
add_bullet(
    doc,
    "Built and trained deep learning models to optimise training loss using MLP, "
    "deep residual networks and CNN architectures (Deep Learning)",
)

# Azure Certification
add_experience_header(
    doc,
    "Azure Data Scientist Associate  —  Microsoft Certification",
    "Sept 2023 – Sept 2026",
)

# BSc
add_experience_header(
    doc,
    "BSc. Engineering (Chemical)  —  Kwame Nkrumah University of Science & Technology",
    "Sept 2015 – July 2019",
)

# ── SECTION 3: RESEARCH & ACADEMIC PROJECTS ───────────────────────────────────

add_section_header(doc, "Research & Academic Projects")

# ---- Project 1 ----
proj1_p = doc.add_paragraph()
proj1_p.paragraph_format.space_before = Pt(5)
proj1_p.paragraph_format.space_after = Pt(1)
r = proj1_p.add_run(
    "Improving Natural Language Understanding via Ensemble Debiasing and Data Cartography"
)
set_run_font(r, size_pt=10.5, bold=True)

proj1_meta_p = doc.add_paragraph()
proj1_meta_p.paragraph_format.space_before = Pt(0)
proj1_meta_p.paragraph_format.space_after = Pt(1)
r_meta = proj1_meta_p.add_run(
    "NLP Final Paper  |  github.com/josephine-amponsah/multi-nli-nlu-optimization  |  Dec 2025"
)
set_run_font(r_meta, size_pt=10, italic=True, color=DARK_GREY)

add_bullet(
    doc,
    "Authored research paper investigating dataset artifacts and spurious correlations in "
    "MultiNLI, fine-tuning ELECTRA-small with full premise-hypothesis pairs and "
    "hypothesis-only baselines",
)
add_bullet(
    doc,
    "Implemented and evaluated ensemble debiasing and dataset cartography as debiasing "
    "strategies, achieving up to 4.6 percentage point improvement in robustness on "
    "CheckList behavioral tests",
)
add_bullet(
    doc,
    "Employed BiLSTM hypothesis-only baselines (InferSent methodology, GloVe embeddings) "
    "alongside transformer-based ELECTRA-small models for controlled architecture comparison",
)
add_bullet(
    doc,
    "Used CheckList behavioral testing framework to evaluate linguistic capabilities "
    "including negation, lexical overlap, quantifiers, and world knowledge beyond "
    "held-out accuracy",
)

# ---- Project 2 ----
proj2_p = doc.add_paragraph()
proj2_p.paragraph_format.space_before = Pt(5)
proj2_p.paragraph_format.space_after = Pt(1)
r = proj2_p.add_run("Neural Network Classifier with Loss Optimisation")
set_run_font(r, size_pt=10.5, bold=True)

proj2_meta_p = doc.add_paragraph()
proj2_meta_p.paragraph_format.space_before = Pt(0)
proj2_meta_p.paragraph_format.space_after = Pt(1)
r_meta = proj2_meta_p.add_run(
    "Deep Learning  |  github.com/josephine-amponsah/nnclassifier-loss-optimization  |  Mar 2026"
)
set_run_font(r_meta, size_pt=10, italic=True, color=DARK_GREY)

add_bullet(
    doc,
    "Designed and trained deep classification networks (Linear, MLP, Deep MLP, Deep Residual "
    "MLP) on image classification tasks using PyTorch",
)
add_bullet(
    doc,
    "Implemented residual connections to mitigate vanishing gradient problems in networks "
    "with 4+ layers, achieving >80% validation accuracy",
)
add_bullet(
    doc,
    "Experimented with classification loss functions (softmax log-likelihood), learning rate "
    "scheduling, and hyperparameter tuning (batch size, hidden dimension, number of layers)",
)
add_bullet(
    doc,
    "Logged training dynamics with TensorBoard to monitor loss and accuracy across epochs",
)

# ---- Project 3 ----
proj3_p = doc.add_paragraph()
proj3_p.paragraph_format.space_before = Pt(5)
proj3_p.paragraph_format.space_after = Pt(1)
r = proj3_p.add_run("Transformer-Based NLP Models")
set_run_font(r, size_pt=10.5, bold=True)

proj3_meta_p = doc.add_paragraph()
proj3_meta_p.paragraph_format.space_before = Pt(0)
proj3_meta_p.paragraph_format.space_after = Pt(1)
r_meta = proj3_meta_p.add_run(
    "NLP  |  github.com/josephine-amponsah/nlp-transformer-models  |  Mar 2026"
)
set_run_font(r_meta, size_pt=10, italic=True, color=DARK_GREY)

add_bullet(
    doc,
    "Implemented and fine-tuned transformer architectures (BERT, ELECTRA) for downstream "
    "NLP classification tasks",
)
add_bullet(
    doc,
    "Explored attention mechanisms, tokenization strategies, and fine-tuning best practices "
    "on sequence classification benchmarks",
)

# ---- Project 4 ----
proj4_p = doc.add_paragraph()
proj4_p.paragraph_format.space_before = Pt(5)
proj4_p.paragraph_format.space_after = Pt(1)
r = proj4_p.add_run("Neural Network Sentiment Classifier")
set_run_font(r, size_pt=10.5, bold=True)

proj4_meta_p = doc.add_paragraph()
proj4_meta_p.paragraph_format.space_before = Pt(0)
proj4_meta_p.paragraph_format.space_after = Pt(1)
r_meta = proj4_meta_p.add_run(
    "NLP  |  github.com/josephine-amponsah/neural-network-sentiment-classifier  |  Mar 2026"
)
set_run_font(r_meta, size_pt=10, italic=True, color=DARK_GREY)

add_bullet(
    doc,
    "Built neural sentiment classification models with optimised architectures (RNNs, "
    "feedforward networks) for fine-grained sentiment analysis",
)
add_bullet(
    doc,
    "Implemented custom optimization routines and model evaluation pipelines with "
    "per-class precision, recall, and F1 metrics",
)

# ---- Project 5 ----
proj5_p = doc.add_paragraph()
proj5_p.paragraph_format.space_before = Pt(5)
proj5_p.paragraph_format.space_after = Pt(1)
r = proj5_p.add_run("Natural Language Preprocessing Optimisation")
set_run_font(r, size_pt=10.5, bold=True)

proj5_meta_p = doc.add_paragraph()
proj5_meta_p.paragraph_format.space_before = Pt(0)
proj5_meta_p.paragraph_format.space_after = Pt(1)
r_meta = proj5_meta_p.add_run(
    "NLP  |  github.com/josephine-amponsah/nl-preprocessing-optimization  |  Mar 2026"
)
set_run_font(r_meta, size_pt=10, italic=True, color=DARK_GREY)

add_bullet(
    doc,
    "Investigated the impact of different preprocessing strategies on downstream sentiment "
    "classification performance",
)
add_bullet(
    doc,
    "Compared tokenization schemes and preprocessing pipelines across multiple model "
    "architectures",
)

# ---- Project 6 ----
proj6_p = doc.add_paragraph()
proj6_p.paragraph_format.space_before = Pt(5)
proj6_p.paragraph_format.space_after = Pt(1)
r = proj6_p.add_run("LLM Fact-Checking Pipeline")
set_run_font(r, size_pt=10.5, bold=True)

proj6_meta_p = doc.add_paragraph()
proj6_meta_p.paragraph_format.space_before = Pt(0)
proj6_meta_p.paragraph_format.space_after = Pt(1)
r_meta = proj6_meta_p.add_run(
    "NLP  |  github.com/josephine-amponsah/factchecking-llm-responses  |  Mar 2026"
)
set_run_font(r_meta, size_pt=10, italic=True, color=DARK_GREY)

add_bullet(
    doc,
    "Developed a fact-checking pipeline to assess and validate LLM-generated responses "
    "against grounding sources",
)
add_bullet(
    doc,
    "Relevant to financial document verification and validation use cases "
    "(e.g., FinBERT-style applications)",
)

# ---- Project 7 ----
proj7_p = doc.add_paragraph()
proj7_p.paragraph_format.space_before = Pt(5)
proj7_p.paragraph_format.space_after = Pt(1)
r = proj7_p.add_run("Aspect-Based Sentiment Analysis with BERT (Prior Project)")
set_run_font(r, size_pt=10.5, bold=True)

proj7_meta_p = doc.add_paragraph()
proj7_meta_p.paragraph_format.space_before = Pt(0)
proj7_meta_p.paragraph_format.space_after = Pt(1)
r_meta = proj7_meta_p.add_run(
    "github.com/josephine-amponsah/absa-based-recommender"
)
set_run_font(r_meta, size_pt=10, italic=True, color=DARK_GREY)

add_bullet(
    doc,
    "Built aspect-based sentiment analysis model using Word2vec and BERT-based ABSA on "
    "user review data to generate fine-grained customer sentiments on specific product aspects",
)
add_bullet(
    doc,
    "Integrated sentiment outputs into a recommender system pipeline",
)

# ── SECTION 4: EXPERIENCE ─────────────────────────────────────────────────────

add_section_header(doc, "Experience")

# Fido
add_experience_header(
    doc,
    "Data Analyst  —  Fido Microcredit (FinTech)",
    "July 2024 – Present",
)

add_italic_label(doc, "Statistical Analysis & ML Research")

add_bullet(
    doc,
    "Optimized statistical models behind fraud monitoring logics using a combination of "
    "central tendency measures and feature engineering techniques, reducing false fraud "
    "flags by 28%",
)
add_bullet(
    doc,
    "Built a new fraud rule supervised learning model to automate the flagging of newly "
    "discovered fraud patterns, which previously required blanket blocking of geohashes "
    "at precision level 5",
)
add_bullet(
    doc,
    "Collaborated with the engineering team on evaluation, validation, and deployment of "
    "new fraud rules to production",
)
add_bullet(
    doc,
    "Analysed the impact of new KYC requirements and flows on risk performance by measuring "
    "defaults across A/B testing cohorts, optimising the intended impact of new feature "
    "implementations",
)
add_bullet(
    doc,
    "Researched and segmented users based on loan racing behaviour and quantified risk "
    "exposure from this cohort, contributing to features implemented in new versions of "
    "risk behavioural models",
)

# WeTheBrands
add_experience_header(
    doc,
    "Data Analyst  —  WeTheBrands (E-Commerce), Germany",
    "Jan 2024 – Nov 2024",
)

# Broadspectrum
add_experience_header(
    doc,
    "Data Analyst  —  Broadspectrum Digital Payments (FinTech)",
    "Jan 2022 – Jan 2024",
)

# ── SECTION 5: SKILLS ─────────────────────────────────────────────────────────

add_section_header(doc, "Skills")

skills_data = [
    (
        "Domains: ",
        "Supervised Learning, Unsupervised Learning, NLP & Transformers, Deep Learning, "
        "Statistical Analysis, Forecasting, Financial Risk Modelling, Data Visualisation, "
        "Recommender Systems",
    ),
    (
        "Frameworks/Tools: ",
        "PyTorch, TensorFlow/Keras, HuggingFace Transformers, Pandas, NumPy, Matplotlib, "
        "Scikit-learn, NLTK, Gensim, FastAPI, Flask, Docker, BigQuery, DBT, PowerBI, "
        "Tableau, GitHub",
    ),
    (
        "Languages: ",
        "Python, R, SQL, DAX, Azure, AWS",
    ),
]

for bold_part, normal_part in skills_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    inline_bold_normal(p, bold_part, normal_part)

# ── SECTION 6: PUBLICATIONS & PRESENTATIONS ───────────────────────────────────

add_section_header(doc, "Publications & Presentations")

pub_entries = [
    (
        'Paper: "Improving Natural Language Understanding via Ensemble Debiasing and Data '
        'Cartography"',
        " — NLP Final Project Paper, University of Texas Austin, 2025. "
        "GitHub: github.com/josephine-amponsah/multi-nli-nlu-optimization",
    ),
    (
        "Talk: ",
        "Statistical & ML-based Fraud Risk Prediction Methods, DevX 2025 Developer Conference",
    ),
    (
        "Talk: ",
        "Geospatial Segmentation for Risk Assessment, PyCon Ghana 2025",
    ),
    (
        "Writing: ",
        "Researching and publishing articles on practical implementations of mathematical, "
        "machine learning, and finance concepts for intermediate Data Analyst & Scientist "
        "audiences",
    ),
]

for bold_part, normal_part in pub_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    inline_bold_normal(p, bold_part, normal_part)

# ── Save ───────────────────────────────────────────────────────────────────────

output_path = "Academic_CV_Josephine_Amponsah_Baah.docx"
doc.save(output_path)
print(f"CV saved to {output_path}")
